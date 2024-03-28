#!/usr/bin/python3

"""This module contains a script that scrape all projects in the alx intranet
    account.
"""
from selenium import webdriver
import requests
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
import webbrowser
from bs4 import BeautifulSoup
import pyperclip
from docx import Document
import os
from selenium.common.exceptions import WebDriverException

browser = webdriver.Chrome()
browser.get("https://intranet.alxswe.com/projects/current")
try:
    # get the email input and input the email
    email_input = browser.find_element(By.ID, "user_email")
    email_input.send_keys(input("Enter you email: ").strip())
    print("email accepted.....")
    
    # get the password input and input the password
    pass_input = browser.find_element(By.ID, "user_password")
    pass_input.send_keys(input("Enter you password: ").strip())
    pass_input.submit()
    print("signing in......")
    print("Welcome to your intranet account")

    time.sleep(2)
    # WebDriverWait(browser, 10).until(EC.presence_of_all_elements_located((By.CLASS_NAME, "btn-default")))
    button = browser.find_element(By.CLASS_NAME, "btn-default")
    # for button in buttons:
    #     if button.text == "Expand all":
    button.click()

    WebDriverWait(browser, 10).until(EC.presence_of_all_elements_located((By.CLASS_NAME, "panel-default")))
    panel_defaults = browser.find_elements(By.CLASS_NAME, "panel-default")
    
    print("\nlooping through panel defaults\n")
    for panel_default in panel_defaults[:5]:
        panel_title = panel_default.find_element(By.CSS_SELECTOR, ".panel-title a")
        print(panel_title.text.upper())
        
        # Create a directory with the project title
        project_directory = panel_title.text.upper().split("\n")[0].replace(" ", "_")
        os.makedirs(project_directory, exist_ok=True)

        list_items = panel_default.find_elements(By.CLASS_NAME, "list-group-item")
        for list_item in list_items[:5]:
            panel_project = list_item.find_element(By.TAG_NAME, "a")
            project_name = panel_project.text
            print(f"\t{project_name}")
            project_url = panel_project.get_attribute("href")
            browser.execute_script(f"window.open('{project_url}', '_blank');")
            browser.switch_to.window(browser.window_handles[1])
            
            page_text = browser.find_element(By.TAG_NAME, "body").text
            links = browser.find_elements(By.CSS_SELECTOR, "#curriculum_navigation_content a")
            resources = []
            for link in links:
                resources_link = link.get_attribute("href")
                if resources_link:
                    browser.execute_script(f"window.open('{resources_link}', '_blank');")
                    browser.switch_to.window(browser.window_handles[2])
                    try:
                        original_url = browser.current_url
                    except Exception:
                        print("Unable to get the url of the resouces")
                    resources.append(original_url)
                    browser.close()
                    browser.switch_to.window(browser.window_handles[1])
            # Save the text to a Word document
            doc = Document()
            heading = doc.add_heading(level=1)
            heading_run = heading.add_run("Resources To Complete This Project")
            heading_run.font.underline = True

            for resource in resources:
                doc.add_paragraph(resource)

                # Add a line break between resources for readability
                doc.add_paragraph("")

            # Add a line break between resource list and project content
            doc.add_paragraph("")

            # Add a centered heading for the project content
            doc.add_heading('PROJECT', level=1).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add the project content
            doc.add_paragraph(page_text)

            # Save the document in the project directory
            doc_file = os.path.join(project_directory, f"{project_name}.docx")
            doc.save(doc_file)

            print(f"Content saved to {project_name}.docx")

            browser.close()
            browser.switch_to.window(browser.window_handles[0])
            # pyperclip.copy(page)
            # webbrowser.open(panel_project.get_attribute("href"))
    

except WebDriverException as e:
    if "rejected by interface blink.mojom.WidgetHost" in str(e):
        print("Handling WidgetHost error...")
        # Add handling code specific to the WidgetHost error here
    else:
        print("An exception occurred:")
        print(e)
time.sleep(5)
browser.quit()  # Close the browser when done


